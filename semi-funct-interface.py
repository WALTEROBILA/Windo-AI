import os
from dotenv import load_dotenv
from langchain_groq import ChatGroq
from typing import TypedDict, Annotated, Optional, List, Literal
from langgraph.graph import StateGraph, add_messages, START, END, MessagesState
import requests
from bs4 import BeautifulSoup
from langgraph.types import Command
from langchain_core.messages import HumanMessage, AIMessage 
from pydantic import BaseModel, Field, field_validator, AliasChoices
from IPython.display import Image, display 
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from bs4 import BeautifulSoup
import time 
from langchain_google_genai import ChatGoogleGenerativeAI
import json
from pypdf import PdfReader
from docx import Document
from datetime import datetime
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import streamlit as st
from io import BytesIO 

st.subheader("Tailored, ATS-friendly Resume and Cover Letter Generator")

load_dotenv()

model = ChatGroq(model="llama-3.3-70b-versatile")

# model = ChatGroq(model="llama-3.1-8b-instant")

# url = "https://reliefweb.int/job/4213251/tupande-mel-data-senior-supervisor"

url = st.text_input("Paste a valid job advertisment url")


## Function to extract the contents of the job advertisement webpage
def read_webpage(url):
    options = Options()

    # makes browser look more real
    options.add_argument("--disable-blink-features=AutomationControlled")
    options.add_argument(
        "user-agent=Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    )


    driver = webdriver.Chrome(options=options)

    #hiding selenium traces
    driver.execute_script("""
        Object.defineProperty(navigator, 'webdriver', {
            get: () => undefined
        })
    """)

    driver.get(url)

    #waiting for JS content
    time.sleep(5)

    html = driver.page_source

    driver.quit()

    soup = BeautifulSoup(html, "html.parser")

    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()

    text = soup.get_text(separator="\n", strip=True)

    return text


content = read_webpage(url)
#print(content)


import os
import json

from PyPDF2 import PdfReader
from docx import Document

from langchain_core.messages import HumanMessage

#Function to extract text from PDF / DOCX documents

def extract_text_from_document(file_path):
    """
    Extracts text from either a .pdf or .docx document.
    """

    if not os.path.exists(file_path):
        return "Error: File not found."

    file_extension = os.path.splitext(file_path)[1].lower()

    #PDF
    if file_extension == ".pdf":

        reader = PdfReader(file_path)

        text = ""

        for page in reader.pages:

            extracted_text = page.extract_text()

            if extracted_text:
                text += extracted_text + "\n"

        return text

    #DOCX
    elif file_extension == ".docx":

        doc = Document(file_path)

        text_list = [
            p.text for p in doc.paragraphs
            if p.text.strip()
        ]

        return "\n".join(text_list)

    #Unsupported
    else:
        return f"Error: Unsupported file format '{file_extension}'."


#Function to extract candidate profile from cover letter

def extract_candidate_profile(document_text):
    """
    Extracts structured candidate information from
    a resume, cover letter, CV, or professional summary.
    """

    prompt = f"""
You are an expert resume parser.

Extract the candidate's information from the text below.

Return ONLY valid JSON.

Use this exact structure:

{{
    "name": "",
    "professional_summary": "",
    "skills": [],
    "education": [],
    "experience": [],
    "certifications": [],
    "projects": [],
    "tools_technologies": [],
    "languages": [],
    "achievements": [],
    "email": "",
    "phone": ""
}}

Rules:
- Return ONLY JSON
- No markdown
- No explanations
- Do not hallucinate information
- Use empty arrays if unavailable

TEXT:
{document_text[:12000]}
"""

    response = model.invoke([
        HumanMessage(content=prompt)
    ])
    
    content = response.content.strip()

    try:

        parsed_data = json.loads(content)

        print("Candidate profile extracted successfully.")

        return parsed_data

    except Exception as e:

        print("JSON parsing failed.")
        print(content)

        raise e 


#Function to extract contents of a resume
def extract_resume_profile(resume_text):

    prompt = f"""
You are an expert resume parser.

Extract information from this resume.

Return ONLY valid JSON.

{{
    "name": "",
    "email": "",
    "phone": "",
    "location": "",
    "professional_summary": "",
    "skills": [],
    "experience": [],
    "education": [
        {{
            "institution":"",
            "degree":"",
            "field":"",
            "duration":"",
            "grade":""
        }}
    ],
    "certifications": [],
    "projects": [],
    "tools_technologies": [],
    "languages": [],
    "achievements": [],
    "references": [
        {{
            "name":"",
            "title":"",
            "organization":"",
            "phone":"",
            "email":""
        }}
    ]    
}}

Rules:
- Return ONLY JSON
- No markdown
- No explanations
- Do not hallucinate information
- Use empty arrays if unavailable
- Extract references if present
- Each reference should include:
    - name
    - title
    - organization
    - phone
    - email
- If references are not present, return:
    "references":[]

EDUCATION

For every education entry extract:
    - institution
    - degree
    - field (if available)
    - duration (e.g. Aug 2018 – Apr 2023)
    - grade/classification/GPA if present

Examples of grade:
    - First Class Honours
    - Second Class Upper
    - Second Class Lower
    - Distinction
    - Merit
    - GPA: 3.78/4.00

If a grade is not provided, return:

"grade":""

RESUME:
{resume_text[:12000]}
"""

    response = model.invoke([
        HumanMessage(content=prompt)
    ])

    content = response.content.strip()

    #Remove markdown code fences if the model adds them
    content = re.sub(
        r"^```(?:json)?",
        "",
        content,
        flags=re.IGNORECASE
    )

    content = re.sub(
        r"```$",
        "",
        content
    )

    content = content.strip()

    try:

        parsed_data = json.loads(content)

        print("Resume profile extracted successfully.")

        return parsed_data

    except Exception as e:

        print("JSON parsing failed.")

        print("\nRAW MODEL OUTPUT:\n")
        print(content)

        return {}

import tempfile 

def save_uploaded_file(uploaded_file):
    """
    Saves a Streamlit uploaded file to a temporary location and returns the file path
    """

    suffix = os.path.splitext(uploaded_file.name)[1]

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp_file:
        tmp_file.write(uploaded_file.getbuffer())
        return tmp_file.name  

#Upload files
cover_letter_file = st.file_uploader("Cover Letter", type=["pdf","docx"])
resume_file = st.file_uploader("Resume", type=["pdf","docx"])

if cover_letter_file and resume_file:

    cover_letter_path = save_uploaded_file(cover_letter_file)
    resume_path = save_uploaded_file(resume_file)

    text = extract_text_from_document(cover_letter_path)
    candidate_profile = extract_candidate_profile(text)

    resume_text = extract_text_from_document(resume_path)
    resume_profile = extract_resume_profile(resume_text)

    st.success("Documents processed succesfully!")

#usage
# if __name__ == "__main__":

#     file_path = r"C:\Users\Admin\Desktop\Windo\Dan Church Aid\Walter Obila Cover Letter.pdf"

#     # Step 1: Extract raw text
#     text = extract_text_from_document(file_path)

#     # Step 2: Extract structured profile
#     candidate_profile = extract_candidate_profile(text)

#     # Step 3: Print result
#     print(json.dumps(candidate_profile, indent=4))


#     # resume profile
#     resume_path = r"C:\Users\Admin\Desktop\Windo\Dalberg Data Insight\ADERO WALTER OBILA CV.pdf"
#     resume_text = extract_text_from_document(resume_path)
#     print(resume_text) 
#     resume_profile = extract_resume_profile(resume_text)



def is_job_advertisement_node(content):
    prompt = f"""
You are a classifier.

Determine whether the following webpage content is a job advertisement/posting.

Respond ONLY with:
YES
or
NO

Webpage content:
{content}
"""

    response = model.invoke([
        HumanMessage(content=prompt)
    ])

    result = response.content.strip().upper()

    if result != "YES":
        st.error("THIS IS NOT A VALID JOB POSTING!")
        st.stop()
    
    print("Advertisment checked and verified as valid")

    return {
        "job_advertisement": "YES"
    }

#Graph's State
class GraphState(TypedDict):
    content:str
    job_advertisement: str

    #outputs
    cover_letter:str
    resume: str

    #cover letter source profile
    candidate_profile: dict

    #resume source profile
    resume_profile: dict

    #raw extracted documents
    source_cover_letter_text: str
    source_resume_text: str


    applicant_header: str
    company_header: str

    resume_summary: str
    resume_core_competencies: list
    resume_experience: list
    resume_projects: list
    resume_education: list
    resume_certifications: list

    resume_references: list

from datetime import datetime
from langchain_core.messages import HumanMessage
import json


#Function to build the cover letter
def cover_letter_builder_node(state: GraphState):

    print("Building the Cover Letter")

    try:

        # -----------------------------------
        # Candidate Profile
        # -----------------------------------
        profile = state["candidate_profile"]

        job_posting = state["content"]

        name = profile.get("name", "")
        email = profile.get("email", "")
        phone = profile.get("phone", "")

        current_date = datetime.now().strftime("%d %B %Y")

        # -----------------------------------
        # Applicant Header
        # -----------------------------------
        applicant_header = f"""
{name}
{email}
{phone}
{current_date}
""".strip()

        #STEP 1:Extract company details using LLM
        extraction_prompt = f"""
Extract the company information from this job posting.

Return ONLY valid JSON in this exact format:

{{
    "company_name": "",
    "company_location": ""
}}

Rules:
- company_location should be formatted like:
  Nairobi, Kenya

- If missing, infer from the posting where possible
- Return ONLY JSON

JOB POSTING:
{job_posting[:4000]}
"""

        extraction_response = model.invoke([
            HumanMessage(content=extraction_prompt)
        ])

        extraction_content = extraction_response.content.strip()

        try:

            extracted_company = json.loads(
                extraction_content
            )

        except Exception:

            extracted_company = {
                "company_name": "",
                "company_location": ""
            }

        company_name = extracted_company.get(
            "company_name",
            ""
        )

        company_location = extracted_company.get(
            "company_location",
            ""
        )

        company_header = f"""
{company_name}
{company_location}
""".strip()

        
        #STEP 2:Generate Cover Letter BODY ONLY
        prompt = f"""
Generate ONLY the BODY of a professional cover letter.

DO NOT include:
- applicant header
- company header
- date
- salutation
- signoff

Include ONLY:
- introductory paragraph
- three body paragraphs
- short closing paragraph

RULES:
- concise and professional
- ATS-friendly
- natural and human sounding
- do not copy the job posting directly
- naturally use keywords
- phrase qualifications as achievements
- never mention lacking qualifications
- mention attraction to the company and role
- avoid placeholders

CANDIDATE PROFILE:
{profile}

JOB POSTING:
{job_posting[:6000]}
"""

        response = model.invoke([
            HumanMessage(content=prompt)
        ])

        cover_letter = response.content.strip()

        print("Cover Letter Generated Successfully")

        return {
            "cover_letter": cover_letter,
            "applicant_header": applicant_header,
            "company_header": company_header
        }

    except Exception as e:

        print("ERROR GENERATING COVER LETTER:")
        print(e)

        return {
            "cover_letter": "Failed to generate cover letter.",
            "applicant_header": "",
            "company_header": ""
        }


from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

#Function to save the cover letter to a word document
def save_cover_letter_to_word(
    applicant_header,
    company_header,
    cover_letter,
    applicant_name,
    file_name="cover_letter.docx"
):

    document = Document()

    #Global Styling
    style = document.styles["Normal"]

    style.font.name = "Calibri"
    style.font.size = Pt(11)

    #Applicant Header (RIGHT)
    sender = document.add_paragraph()

    sender.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    sender.add_run(applicant_header)

    sender.paragraph_format.space_after = Pt(12)

    #Company Header (LEFT)
    company = document.add_paragraph()

    company.alignment = WD_ALIGN_PARAGRAPH.LEFT

    company.add_run(company_header)

    company.paragraph_format.space_after = Pt(12)

    #Salutation
    salutation = document.add_paragraph()

    salutation.alignment = WD_ALIGN_PARAGRAPH.LEFT

    salutation.add_run("Dear Hiring Manager,")

    salutation.paragraph_format.space_after = Pt(12)

    #Body Paragraphs
    paragraphs = [
        p.strip()
        for p in cover_letter.splitlines()
        if p.strip()
    ]

    for para in paragraphs:

        p = document.add_paragraph(para)

        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        p.paragraph_format.space_after = Pt(10)

    #Sign Off
    closing = document.add_paragraph()

    closing.alignment = WD_ALIGN_PARAGRAPH.LEFT

    closing.add_run("Sincerely,")

    closing.paragraph_format.space_after = Pt(0)

    #Applicant Name
    signature = document.add_paragraph()

    signature.alignment = WD_ALIGN_PARAGRAPH.LEFT

    signature.add_run(applicant_name)

    signature.paragraph_format.space_before = Pt(0)

    #Save
    document.save(file_name)

    print(f"Saved as {file_name}")

    # Save to memory instead of disk
    output = BytesIO()
    document.save(output)
    output.seek(0)

    return output


#Function to build the resume
def resume_builder_node(state: GraphState):

    print("Building Resume")

    profile = state["resume_profile"]

    job_posting = state["content"]

    prompt = f"""
You are an expert ATS resume writer.

Optimize the candidate's resume for the job posting.

Return ONLY valid JSON.

{{
    "professional_summary": "",
    "core_competencies": [],
    "experience": [
        {{
            "title":"",
            "organization":"",
            "duration":"",
            "responsibilities":[]
        }}
    ],
    "projects": [
        {{
            "title":"",
            "description":"",
            "technologies":[]
        }}
    ],
    "education": [
        {{
            "institution":"",
            "degree":"",
            "field":"",
            "duration":"",
            "grade":""
        }}
    ],
    "certifications": [],
    "references":[]
}}

Rules:

- Use the information provided for most sections of the resume.
- Tweak the contents of the professional summary in order to match the resume profile to the job description. {job_posting[:7000]}. Do not mention the organization being applied to, make it seem like a generic summary, but with the role being applied for aligning with the contents of the professional summary.

- For each experience entry preserve:
    - title
    - organization
    - duration
    - responsibilities

    Responsibilities should remain as a list of bullet points. Do not remove or collapse them.

- Rewrite experience bullets to align with the role. Ensure that the points in the bullet align with the role. 
- Three to four bullets per company worked for. 
- Ensure the bullets are significant, heavy showing experience and prowess. 
- Ensure you quantify some of the output where possible, depicting value brought numerically.
- Include important keywords naturally.
- You are allowed to embellish responsibilities to fit with the role being advertised. 
- Ensure the applicant comes across as qualified and experienced.

- ATS friendly.

- For the core competencies, do not list all skills lifted from the resume. Select the ones relevant to the job posting, ensuring a majority of those mentioned in the posting are included. I want 
    the core competencies to be grouped into categories if applying for tech roles, by competency group and the skills/techniologies entailed. e.g
        - Data Analysis - R, SPSS, Excel
        - Machine Learning Modelling - Scikit-learn, TensorFlow
        - Data Visualization - PowerBI, Tableau
- Ensure you have listed relevant competencies, as wells as the skills involved.
- The core competencies section can't be empty, can't be generic. Ensure it is convincingly populated.

EDUCATION
- Preserve all education entries.
- For each education entry preserve:
    - institution
    - degree
    - field
    - duration
    - grade
- Do not invent grades.
- Do not remove classifications such as:
    - First Class Honours
    - Second Class Upper
    - Distinction
    - Merit
    - GPA
- Return education using this structure:
    "education":[
        {{
            "institution":"",
            "degree":"",
            "field":"",
            "duration":"",
            "grade":""
        }}
]

PROJECTS:
- Include ONLY projects that are relevant to the target role.
- If none of the projects are relevant, return:
    "projects": []

- Do not invent projects.
- Preserve factual accuracy.
- Rewrite project descriptions to emphasize skills, methodologies, tools and technologies that align with the job posting.

- Prioritize:
    - software tools
    - programming languages
    - frameworks
    - analytical methods
    - research methods
    - technologies
that are most relevant to the job posting.

- Preserve the original technologies used.
- Do not add technologies that were never used.

Return projects using this structure:

"projects": [
    {{
        "title":"",
        "description":"",
        "technologies":[]
    }}
]

- Return ONLY JSON.
- Preserve references exactly as provided
- Do not invent references

Candidate Profile:
{json.dumps(profile, indent=2)}

Job Posting:
{job_posting[:7000]}
"""

    response = model.invoke([
        HumanMessage(content=prompt)
    ])

    content = response.content.strip()

    content = content.replace("```json", "")
    content = content.replace("```", "")
    content = content.strip()

    resume_data = json.loads(content)

    print("Resume Generated Successfully")

    print(json.dumps(
    resume_data,
    indent=4
    ))

    return {
        "resume_summary": resume_data.get(
            "professional_summary",
            ""
        ),
        "resume_core_competencies": resume_data.get(
            "core_competencies",
            []
        ),
        "resume_experience": resume_data.get(
            "experience",
            []
        ),
        "resume_projects": resume_data.get(
            "projects",
            []
        ),
        "resume_education": resume_data.get(
            "education",
            []
        ),
        "resume_certifications": resume_data.get(
            "certifications",
            []
        ),
        "resume_references": resume_data.get(
            "references",
            []
        )
    }
    
#Function to save the resume
def save_resume_to_word(
    resume_profile,
    summary,
    competencies,
    experience,
    projects,
    education,
    certifications,
    references,
    file_name="resume.docx"
):

    document = Document()

    style = document.styles["Normal"]

    style.font.name = "Calibri"
    style.font.size = Pt(11)

    name = resume_profile.get("name", "")
    email = resume_profile.get("email", "")
    phone = resume_profile.get("phone", "")

    #Resume header
    header = document.add_paragraph()

    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    name_run = header.add_run(
        f"{name.upper()}\n"
    )

    name_run.bold = True
    name_run.font.size = Pt(16)

    contact_run = header.add_run(
        f"{email} | {phone}"
    )

    contact_run.font.size = Pt(11)

    header.paragraph_format.space_after = Pt(12)

    #Professional summary
    document.add_heading(
        "PROFESSIONAL SUMMARY",
        level=1
    )
    document.add_paragraph(summary)

    #Core competencies
    document.add_heading(
        "CORE COMPETENCIES",
        level=1
    )

    for skill in competencies:
        document.add_paragraph(
            skill,
            style="List Bullet"
        )

    #Professional Experience
    document.add_heading(
        "PROFESSIONAL EXPERIENCE",
        level=1
    )

    for job in experience:

        title = job.get("title", "")

        organization = job.get(
            "organization",
            ""
        )

        duration = job.get(
            "duration",
            ""
        )

        # Job title
        p =document.add_paragraph()

        p.paragraph_format.space_after = Pt(0)

        run = p.add_run(
            f"{title}|{organization}"
        )

        run.bold = True

        #Duration
        if duration:
            d = document.add_paragraph(duration)

            d.paragraph_format.space_before = Pt(0)
            d.paragraph_format.space_after = Pt(2)

        # Responsibilities
        for bullet in job.get(
            "responsibilities",
            []
        ):
            
            document.add_paragraph(
                bullet,
                style="List Bullet"
            )


    #Featured projects
    if projects:

        document.add_heading(
            "FEATURED PROJECTS",
            level=1
        )

        for project in projects:

            title = project.get(
                "title",
                ""
            )

            description = project.get(
                "description",
                ""
            )

            technologies = project.get(
                "technologies",
                []
            )

            # Project title
            p = document.add_paragraph()

            p.paragraph_format.space_after = Pt(0)

            run = p.add_run(title)

            run.bold = True

            # Description
            if description:

                desc = document.add_paragraph(
                    description
                )

                desc.paragraph_format.space_after = Pt(2)

            # Technologies
            if technologies:

                tech = document.add_paragraph()

                tech.paragraph_format.space_after = Pt(6)

                tech_run = tech.add_run(
                    "Technologies: "
                    + ", ".join(technologies)
                )

                tech_run.italic = True

    #Education
    document.add_heading(
        "EDUCATION",
        level=1
    )

    for edu in education:

        degree = edu.get("degree", "")
        field = edu.get("field", "")
        institution = edu.get("institution", "")
        duration = edu.get("duration", "")
        grade = edu.get("grade", "")

        #Degree + field + institution
        p = document.add_paragraph()

        p.paragraph_format.space_after = Pt(0)

        display_degree = degree

        if field and field.lower() not in degree.lower():
            display_degree += f" in {field}"

        run = p.add_run(
            f"{display_degree} | {institution}"
        )

        run.bold = True

        #Duration
        if duration:

            d = document.add_paragraph(duration)

            d.paragraph_format.space_before = Pt(0)
            d.paragraph_format.space_after = Pt(0)

        #Grade/Classification
        if grade:

            g = document.add_paragraph(grade)

            g.paragraph_format.space_before = Pt(0)
            g.paragraph_format.space_after = Pt(6)

            g.runs[0].italic = True

    #Certifications
    if certifications:

        document.add_heading(
            "CERTIFICATIONS",
            level=1
        )

        for cert in certifications:

            if isinstance(cert, dict):

                name = cert.get("name", "")
                issuer = cert.get("issuer", "")
                date = cert.get("date", "")

                p = document.add_paragraph()

                run = p.add_run(name)
                run.bold = True

                if issuer:
                    p.add_run(f"\n{issuer}")

                if date:
                    p.add_run(f"\n{date}")

                p.paragraph_format.space_after = Pt(6)

            else:

                document.add_paragraph(
                    str(cert),
                    style="List Bullet"
                )


#References
    document.add_page_break()

    document.add_heading(
        "REFERENCES",
        level=1
    )

    if references:

        for ref in references:

            name = ref.get("name", "")
            title = ref.get("title", "")
            organization = ref.get("organization", "")
            phone = ref.get("phone", "")
            email = ref.get("email", "")

            p = document.add_paragraph()

            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(6)

            #Name (bold)
            run = p.add_run(name)
            run.bold = False 

            #Remaining details
            if title:
                p.add_run(f"\n{title}")

            if organization:
                p.add_run(f"\n{organization}")

            if phone:
                p.add_run(f"\n{phone}")

            if email:
                p.add_run(f"\n{email}")

    else:

        document.add_paragraph(
            "References available upon request."
        )   

    #Save the resume
    # document.save(file_name)

    # print(f"Saved as {file_name}")
    output = BytesIO()
    document.save(output)
    output.seek(0)

    return output

def advertisement_router(state: GraphState):
    if state["job_advertisement"].strip().upper() == "YES":
        print("Advertisement checked and verified as valid")
        return "yes"
    
    print("THIS IS NOT A VALID JOB POSTING!")

    return "no"



#Building the Graph Structure
graph = StateGraph(GraphState)

graph.add_node("Check Advertisement", is_job_advertisement_node)
graph.add_node("Cover Letter Generator", cover_letter_builder_node)
graph.add_node("Resume Builder", resume_builder_node)

graph.set_entry_point("Check Advertisement")


# graph.add_conditional_edges(
#     "Check Advertisement",
#     advertisement_router,
#     {
#         "yes": "Cover Letter Generator",
#         "no": END
#     }
# )


graph.add_edge("Check Advertisement", "Cover Letter Generator")
graph.add_edge("Cover Letter Generator", "Resume Builder")
graph.add_edge("Resume Builder", END)

app = graph.compile()

ready = (
    content
    and candidate_profile
    and resume_profile
)

generate = st.button(
    "🚀 Generate Cover Letter and Resume",
    type = "primary",
    disabled=not ready
)

if generate:

    with st.spinner("Generating your documents..."):

        result = app.invoke({
            "content": content,
            "candidate_profile": candidate_profile,
            "resume_profile": resume_profile
        })

        cover_letter_doc = save_cover_letter_to_word(
            applicant_header=result["applicant_header"],
            company_header=result["company_header"],
            cover_letter=result["cover_letter"],
            applicant_name=candidate_profile.get("name", "")
        )

        resume_doc = save_resume_to_word(
            resume_profile=result["resume_profile"],
            summary=result["resume_summary"],
            competencies=result["resume_core_competencies"],
            experience=result["resume_experience"],
            projects=result["resume_projects"],
            education=result["resume_education"],
            certifications=result["resume_certifications"],
            references=result["resume_references"]
        )

    st.success("Documents generated successfully!")

    st.download_button(
        label="📄 Download Cover Letter",
        data=cover_letter_doc,
        file_name="Cover_Letter.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    st.download_button(
        label="📄 Download Resume",
        data=resume_doc,
        file_name="Resume.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# result = app.invoke({
#     "content": content,
#     "candidate_profile": candidate_profile,
#     "resume_profile": resume_profile
# })


# cover_letter = result["cover_letter"]


# cover_letter_doc = save_cover_letter_to_word(
#     applicant_header=result["applicant_header"],
#     company_header = result["company_header"],
#     cover_letter=result["cover_letter"],
#     applicant_name=candidate_profile.get("name", "")
#     # file_name="Walter_Obila_Cover_Letter.docx"
# )


# resume_doc = save_resume_to_word(
#     resume_profile=result["resume_profile"],
#     summary=result["resume_summary"],
#     competencies=result["resume_core_competencies"],
#     experience=result["resume_experience"],
#     projects=result["resume_projects"],
#     education=result["resume_education"],
#     certifications=result["resume_certifications"],
#     references=result["resume_references"]
#     # file_name="Walter_Obila_Resume.docx"
# )

# st.download_button(
#     label = "📄 Download Cover Letter",
#     data = cover_letter_doc,
#     file_name = "Cover_Letter.docx",
#     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
# )

# st.download_button(
#     label = "📄 Download Resume",
#     data = resume_doc,
#     file_name = "resume.docx",
#     mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
# )